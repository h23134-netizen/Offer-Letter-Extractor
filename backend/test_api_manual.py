import requests
import os

BASE_URL = "http://localhost:8000"

def test_health():
    try:
        r = requests.get(f"{BASE_URL}/health")
        print(f"Health Check: {r.status_code} - {r.json()}")
    except Exception as e:
        print(f"Health Check Failed: {e}")

def create_dummy_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Offer Letter")
    doc.add_paragraph("Candidate Name: John Doe")
    doc.add_paragraph("Role: Software Engineer")
    doc.add_paragraph("CTC: 12 LPA")
    doc.save("test_offer.docx")
    print("Created test_offer.docx")

def test_parse():
    if not os.path.exists("test_offer.docx"):
        create_dummy_docx()
    
    files = {'files': open('test_offer.docx', 'rb')}
    try:
        r = requests.post(f"{BASE_URL}/parse", files=files)
        print(f"Parse Check: {r.status_code}")
        print(r.json())
    except Exception as e:
        print(f"Parse Check Failed: {e}")

if __name__ == "__main__":
    test_health()
    test_parse()
